import os
import re
import tempfile
from pathlib import Path
from typing import Optional

import httpx
from docx import Document
from fastapi import FastAPI
from openai import OpenAI
from pydantic import BaseModel
from pypdf import PdfReader

app = FastAPI()

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=OPENAI_API_KEY) if OPENAI_API_KEY else None


class ParseAdminGuidanceRequest(BaseModel):
    file_url: str
    mode: str
    admin_prompt: Optional[str] = None


class ParseAdminGuidanceResponse(BaseModel):
    status: str
    file_name: Optional[str] = None
    parsed_text: Optional[str] = None
    file_summary: Optional[str] = None
    parse_error: Optional[str] = None


def clean_extracted_text(text: str) -> str:
    if not text:
        return ""

    text = text.replace("\x00", " ")
    text = text.replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = "\n".join(line.strip() for line in text.splitlines())
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


def truncate_chars(text: str, limit: int = 15000) -> str:
    if not text:
        return ""
    return text[:limit]


async def download_file(file_url: str) -> tuple[str, str]:
    async with httpx.AsyncClient(timeout=60.0, follow_redirects=True) as http:
        resp = await http.get(file_url)
        resp.raise_for_status()

        raw_name = Path(file_url.split("?")[0]).name or "uploaded_file"
        suffix = Path(raw_name).suffix.lower()

        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(resp.content)
            tmp_path = tmp.name

    return tmp_path, raw_name


def parse_docx(local_path: str) -> str:
    doc = Document(local_path)
    parts = []

    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            parts.append(txt)

    for table in doc.tables:
        for row in table.rows:
            row_values = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_values.append(cell_text)
            if row_values:
                parts.append(" | ".join(row_values))

    return "\n".join(parts)


def parse_pdf(local_path: str) -> str:
    reader = PdfReader(local_path)
    parts = []

    for page in reader.pages:
        txt = page.extract_text()
        if txt:
            parts.append(txt)

    return "\n".join(parts)


def parse_txt(local_path: str) -> str:
    with open(local_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def extract_text_from_file(local_path: str, file_name: str) -> str:
    ext = Path(file_name).suffix.lower()

    if ext == ".docx":
        return parse_docx(local_path)
    if ext == ".pdf":
        return parse_pdf(local_path)
    if ext in [".txt", ".md"]:
        return parse_txt(local_path)

    raise ValueError(f"Unsupported file type: {ext}")


def build_fallback_summary(parsed_text: str, mode: str) -> str:
    preview = truncate_chars(parsed_text, 2500)

    return (
        f"Use this uploaded document as approved guidance for {mode} mode. "
        f"Do not quote it verbatim. Extract the preferred response style, coaching logic, "
        f"question types covered, helpful references, and any avoid patterns. "
        f"Base replies on the style and structure implied by this document.\n\n"
        f"Document preview:\n{preview}"
    )


def generate_guidance_summary_with_openai(
    parsed_text: str,
    mode: str,
    admin_prompt: Optional[str]
) -> str:
    if not client:
        return build_fallback_summary(parsed_text, mode)

    trimmed_text = truncate_chars(parsed_text, 12000)
    admin_prompt = admin_prompt or ""

    system_prompt = (
        "You are converting an uploaded admin example document into a compact prompt-guidance block "
        "for an AI assistant. "
        "The document is NOT knowledge-base content and NOT factual retrieval content. "
        "It is approved guidance showing preferred answer style, coaching logic, and routing patterns.\n\n"
        "Return a concise guidance block in plain text.\n"
        "Do not return JSON.\n"
        "Do not quote large sections.\n"
        "Do not copy the document verbatim.\n"
        "Do not include markdown headings.\n"
        "Keep it compact, practical, and prompt-ready.\n\n"
        "The guidance block must capture:\n"
        "- what kinds of user questions this document helps answer\n"
        "- preferred coaching logic\n"
        "- preferred tone/style\n"
        "- which tools/concepts/framework references should be reinforced\n"
        "- what the assistant should avoid\n"
        "- instruction that the assistant should follow the pattern, not copy exact wording\n"
    )

    user_prompt = f"""
Mode: {mode}

Current admin prompt:
{admin_prompt}

Uploaded document text:
{trimmed_text}

Now produce a compact guidance block for the assistant prompt.
"""

    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        temperature=0.2,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
    )

    summary = resp.choices[0].message.content or ""
    return summary.strip()


@app.post("/parse-admin-guidance-file", response_model=ParseAdminGuidanceResponse)
async def parse_admin_guidance_file(payload: ParseAdminGuidanceRequest):
    local_path = None
    file_name = None

    try:
        local_path, file_name = await download_file(payload.file_url)

        extracted_text = extract_text_from_file(local_path, file_name)
        parsed_text = clean_extracted_text(extracted_text)

        if not parsed_text:
            return ParseAdminGuidanceResponse(
                status="error",
                file_name=file_name,
                parsed_text="",
                file_summary="",
                parse_error="No readable text found in uploaded file."
            )

        file_summary = generate_guidance_summary_with_openai(
            parsed_text=parsed_text,
            mode=payload.mode,
            admin_prompt=payload.admin_prompt
        )

        return ParseAdminGuidanceResponse(
            status="success",
            file_name=file_name,
            parsed_text=parsed_text,
            file_summary=file_summary,
            parse_error=""
        )

    except Exception as e:
        return ParseAdminGuidanceResponse(
            status="error",
            file_name=file_name,
            parsed_text="",
            file_summary="",
            parse_error=str(e)
        )

    finally:
        if local_path and os.path.exists(local_path):
            try:
                os.remove(local_path)
            except Exception:
                pass
